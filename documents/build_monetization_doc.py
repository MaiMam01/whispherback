"""
Build WhisperBack_Monetization_Strategy.docx — a branded 3-4 page
market research & GTM strategy document.

Design language:
  - Brand deep purple   #2A2057
  - Brand purple        #4C3A8A
  - Soft purple bg      #F4F1FB
  - Gold accent         #C9A34B
  - Ink                 #15122B
  - Muted               #6B6782
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from copy import deepcopy

# ---- Brand palette --------------------------------------------------
BRAND_DEEP   = "2A2057"
BRAND        = "4C3A8A"
BRAND_2      = "6B56B8"
BRAND_SOFT   = "F4F1FB"
BRAND_SOFT_2 = "EDE7F8"
ACCENT       = "C9A34B"
ACCENT_SOFT  = "F7EFD8"
INK          = "15122B"
INK_2        = "3A354C"
MUTED        = "6B6782"
LINE         = "E6E0F2"
WHITE        = "FFFFFF"


def hex_rgb(h):
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ---- XML helpers ----------------------------------------------------
def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="DDDDDD", size_eighths=4, sides=("top","left","bottom","right")):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for side in sides:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size_eighths))
        b.set(qn("w:color"), color)
        tc_borders.append(b)


def set_cell_margins(cell, top=80, left=140, bottom=80, right=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def paragraph_shading(paragraph, hex_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def paragraph_border(paragraph, hex_color="C9A34B", side="left", size_eighths=24):
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = p_pr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        p_pr.append(pBdr)
    b = OxmlElement(f"w:{side}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size_eighths))
    b.set(qn("w:space"), "8")
    b.set(qn("w:color"), hex_color)
    pBdr.append(b)


def paragraph_spacing(paragraph, before=0, after=0, line=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def styled_run(p, text, *, size=10.5, bold=False, color=INK, font="Segoe UI", italic=False):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = hex_rgb(color)
    rPr = r._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    return r


# ---- Document setup -------------------------------------------------
doc = Document()

# Page margins (tight for a 3-4 pager)
for section in doc.sections:
    section.top_margin = Mm(13)
    section.bottom_margin = Mm(13)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)
    section.header_distance = Mm(6)
    section.footer_distance = Mm(6)

# Default style
style_normal = doc.styles["Normal"]
style_normal.font.name = "Segoe UI"
style_normal.font.size = Pt(10)
style_normal.font.color.rgb = hex_rgb(INK_2)


# =====================================================================
# PAGE 1 — COVER & EXECUTIVE SUMMARY
# =====================================================================

# Top brand bar (single-cell table, deep purple block)
top_bar = doc.add_table(rows=1, cols=2)
top_bar.autofit = False
top_bar.columns[0].width = Cm(13)
top_bar.columns[1].width = Cm(4.4)
cell_l = top_bar.rows[0].cells[0]
cell_r = top_bar.rows[0].cells[1]
for c in (cell_l, cell_r):
    set_cell_shading(c, BRAND_DEEP)
    set_cell_margins(c, top=200, left=200, bottom=200, right=200)
    set_cell_borders(c, color=BRAND_DEEP, size_eighths=2)

# Logo cell
p = cell_l.paragraphs[0]
paragraph_spacing(p, after=2)
styled_run(p, "WHISPERBACK", size=18, bold=True, color=WHITE, font="Segoe UI")
p2 = cell_l.add_paragraph()
paragraph_spacing(p2, after=0)
styled_run(p2, "Your Personalized Audio Whisperer", size=9, color=ACCENT, italic=True)

# Right cell — meta
pr = cell_r.paragraphs[0]
pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
paragraph_spacing(pr, after=2)
styled_run(pr, "FOUS VENTURES", size=8.5, bold=True, color=ACCENT)
pr2 = cell_r.add_paragraph()
pr2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
paragraph_spacing(pr2, after=0)
styled_run(pr2, "Strategy · Discovery", size=8, color=WHITE)

# Eyebrow
p = doc.add_paragraph()
paragraph_spacing(p, before=14, after=2)
styled_run(p, "MARKET RESEARCH  ·  GTM STRATEGY", size=8.5, bold=True, color=ACCENT)

# Title
p = doc.add_paragraph()
paragraph_spacing(p, before=0, after=4)
styled_run(p, "Monetization & Pricing Strategy", size=24, bold=True, color=BRAND_DEEP)

# Subtitle
p = doc.add_paragraph()
paragraph_spacing(p, before=0, after=10, line=1.35)
styled_run(p,
    "Five legitimate revenue streams, the recommended pricing model, "
    "and a 90-day go-to-market plan for WhisperBack — informed by the "
    "personal-development, wellness, and audio-app categories.",
    size=11, color=INK_2)

# Meta row (3 columns)
meta = doc.add_table(rows=1, cols=3)
meta.autofit = False
for col, w in zip(meta.columns, (Cm(5.8), Cm(5.8), Cm(5.8))):
    col.width = w
labels = [("CATEGORY", "Wellness · Productivity · Audio"),
          ("PRIMARY MARKETS", "Pakistan · GCC · India · Global EN"),
          ("DOC", "Strategy · v1.0 · Apr 2026")]
for i, (lbl, val) in enumerate(labels):
    c = meta.rows[0].cells[i]
    set_cell_borders(c, color=ACCENT, size_eighths=8, sides=("top",))
    set_cell_borders(c, color="FFFFFF", size_eighths=2, sides=("left","right","bottom"))
    set_cell_margins(c, top=80, left=0, bottom=40, right=0)
    p = c.paragraphs[0]
    paragraph_spacing(p, after=2)
    styled_run(p, lbl, size=7.5, bold=True, color=MUTED)
    p2 = c.add_paragraph()
    paragraph_spacing(p2, after=0)
    styled_run(p2, val, size=9.5, color=INK)

# Executive summary heading
p = doc.add_paragraph()
paragraph_spacing(p, before=14, after=4)
styled_run(p, "Executive Summary", size=14, bold=True, color=BRAND_DEEP)
# Gold rule
p = doc.add_paragraph()
paragraph_spacing(p, before=0, after=8)
paragraph_border(p, hex_color=ACCENT, side="bottom", size_eighths=12)

# Body paragraph
p = doc.add_paragraph()
paragraph_spacing(p, after=8, line=1.4)
styled_run(p,
    "WhisperBack sits at the intersection of three high-engagement categories — "
    "personal development (affirmations, habit reminders), wellness (sleep, mindfulness), "
    "and learning (language loops, memorisation). These categories pay well: top apps in this "
    "space (Calm, Headspace, ThinkUp, I Am, Motivation) generate revenue overwhelmingly "
    "through subscription, with secondary streams from content packs, advertising, and B2B "
    "licensing.",
    size=10, color=INK_2)
p = doc.add_paragraph()
paragraph_spacing(p, after=10, line=1.4)
styled_run(p,
    "We recommend a ", size=10, color=INK_2)
styled_run(p, "Hybrid Freemium model", size=10, bold=True, color=BRAND_DEEP)
styled_run(p,
    " — free tier monetised by ads, paid tier driven by an annual subscription "
    "with a lifetime escape hatch — supported by four secondary streams that compound over time.",
    size=10, color=INK_2)

# Key takeaways callout
take = doc.add_table(rows=1, cols=1)
take.autofit = False
take.columns[0].width = Cm(17.4)
tc = take.rows[0].cells[0]
set_cell_shading(tc, BRAND_SOFT)
set_cell_borders(tc, color=BRAND, size_eighths=2, sides=("top","right","bottom"))
set_cell_borders(tc, color=BRAND, size_eighths=24, sides=("left",))
set_cell_margins(tc, top=140, left=180, bottom=140, right=180)
p = tc.paragraphs[0]
paragraph_spacing(p, after=4)
styled_run(p, "AT A GLANCE", size=8, bold=True, color=BRAND_DEEP)
bullets = [
    ("Primary revenue", "Tiered subscription (Monthly · Annual · Lifetime)"),
    ("Secondary streams", "AdMob · Content Packs · B2B Licensing · Affiliate"),
    ("Recommended price floor", "$4.99/mo  ·  $39.99/yr  ·  $79.99 Lifetime"),
    ("Free trial", "7 days of Premium on first signup"),
    ("Regional pricing", "40–60% of USD for PK · IN · GCC via store rules"),
]
for label, val in bullets:
    p = tc.add_paragraph()
    paragraph_spacing(p, after=2, line=1.35)
    styled_run(p, "▸  ", size=10, bold=True, color=ACCENT)
    styled_run(p, f"{label}  —  ", size=9.5, bold=True, color=BRAND_DEEP)
    styled_run(p, val, size=9.5, color=INK_2)


# =====================================================================
# PAGE 2 — FIVE MONETIZATION STREAMS
# =====================================================================
add_page_break(doc)

p = doc.add_paragraph()
paragraph_spacing(p, before=0, after=2)
styled_run(p, "01", size=9, bold=True, color=ACCENT)
p = doc.add_paragraph()
paragraph_spacing(p, before=0, after=4)
styled_run(p, "Five Ways to Monetize WhisperBack", size=16, bold=True, color=BRAND_DEEP)
p = doc.add_paragraph()
paragraph_spacing(p, before=0, after=10)
paragraph_border(p, hex_color=ACCENT, side="bottom", size_eighths=12)

streams = [
    {
        "n": "01",
        "title": "Subscription (Premium Tier)",
        "tag": "PRIMARY · 60–75% of revenue",
        "body": "Recurring access to unlimited playlists, cloud sync, multi-device, no ads, "
                "and higher-quality recording. The dominant model for wellness & self-improvement "
                "apps — predictable MRR, compound growth, cleanest investor signal.",
        "stack": "Monthly · Annual (33% saving) · Family Plan (up to 6 seats)",
    },
    {
        "n": "02",
        "title": "In-App Advertising (Google AdMob)",
        "tag": "SUPPORT · floor revenue from free users",
        "body": "AdSense is for websites — on mobile we use AdMob (Google's mobile network). "
                "Non-intrusive banners on low-engagement screens plus rewarded video to temporarily "
                "unlock a feature. Monetises the 95–98% who never subscribe; removed on Premium upgrade — itself a conversion lever.",
        "stack": "Banner · Native · Rewarded Video · Interstitial (used sparingly)",
    },
    {
        "n": "03",
        "title": "One-Time Lifetime Unlock",
        "tag": "ADD-ON · captures subscription-averse users",
        "body": "A single payment that unlocks Premium forever — captures the 5–10% of buyers "
                "who refuse subscriptions outright. Higher upfront cash for early runway, easy to "
                "promote during seasonal campaigns (New Year, Ramadan, back-to-school).",
        "stack": "$79.99 one-time · Launch-window discount $59.99 (first 90 days)",
    },
    {
        "n": "04",
        "title": "Curated Content Packs (In-App Purchases)",
        "tag": "GROWTH · creator-economy upside",
        "body": "Professionally produced audio packs sold inside the app — “Morning Motivation”, "
                "“Sleep Affirmations”, “Quranic Whispers”, “Spoken Urdu Booster”. 70/30 revenue share "
                "with creators turns the app into a marketplace, not just a tool, and the catalog compounds over time.",
        "stack": "$2.99 – $14.99 per pack · Bundles · Seasonal/themed releases",
    },
    {
        "n": "05",
        "title": "B2B Licensing (Coaches, Clinics, Schools)",
        "tag": "ENTERPRISE · sticky high-LTV",
        "body": "A Pro tier for therapists, life coaches, language tutors, religious institutions, "
                "and corporate wellness programs to deliver curated audio to clients & students. Long "
                "contracts, low churn, and a credibility halo for the consumer app.",
        "stack": "$24.99/mo Coach · Custom seats for orgs · White-label add-on",
    },
]

# Render streams in a 2-column grid (3 rows: 2+2+1, last spans both cols)
def render_stream_card(cell, s):
    set_cell_shading(cell, BRAND_SOFT)
    set_cell_margins(cell, top=80, left=140, bottom=80, right=140)
    set_cell_borders(cell, color=BRAND, size_eighths=8, sides=("left",))
    set_cell_borders(cell, color=BRAND_SOFT, size_eighths=2, sides=("top","right","bottom"))

    # Title row: number + title
    p = cell.paragraphs[0]
    paragraph_spacing(p, after=1)
    styled_run(p, s["n"] + "   ", size=10, bold=True, color=ACCENT)
    styled_run(p, s["title"], size=10.5, bold=True, color=BRAND_DEEP)

    p = cell.add_paragraph()
    paragraph_spacing(p, after=2)
    styled_run(p, s["tag"], size=7.5, bold=True, color=ACCENT)

    p = cell.add_paragraph()
    paragraph_spacing(p, after=2, line=1.3)
    styled_run(p, s["body"], size=8.5, color=INK_2)

    p = cell.add_paragraph()
    paragraph_spacing(p, after=0, line=1.3)
    styled_run(p, "Mechanic.  ", size=8, bold=True, color=INK)
    styled_run(p, s["stack"], size=8, italic=True, color=BRAND)


# Build 3-row, 2-col grid; last row merged to span 2 cols
grid = doc.add_table(rows=3, cols=2)
grid.autofit = False
grid.columns[0].width = Cm(8.7)
grid.columns[1].width = Cm(8.7)

# Row 1: streams 0, 1
render_stream_card(grid.rows[0].cells[0], streams[0])
render_stream_card(grid.rows[0].cells[1], streams[1])

# Row 2: streams 2, 3
render_stream_card(grid.rows[1].cells[0], streams[2])
render_stream_card(grid.rows[1].cells[1], streams[3])

# Row 3: stream 4 spans both columns
merged_cell = grid.rows[2].cells[0].merge(grid.rows[2].cells[1])
render_stream_card(merged_cell, streams[4])

# Set column widths explicitly on each cell to avoid autofit weirdness
for row in grid.rows:
    for cell in row.cells:
        for tcW in cell._tc.iter(qn("w:tcW")):
            pass
    if len(row.cells) >= 2 and row.cells[0]._tc != row.cells[1]._tc:
        row.cells[0].width = Cm(8.7)
        row.cells[1].width = Cm(8.7)


# =====================================================================
# PAGE 3 — RECOMMENDED PRICING MODEL
# =====================================================================
add_page_break(doc)

p = doc.add_paragraph()
paragraph_spacing(p, before=0, after=2)
styled_run(p, "02", size=9, bold=True, color=ACCENT)
p = doc.add_paragraph()
paragraph_spacing(p, before=0, after=4)
styled_run(p, "Recommended Pricing Model", size=16, bold=True, color=BRAND_DEEP)
p = doc.add_paragraph()
paragraph_spacing(p, before=0, after=10)
paragraph_border(p, hex_color=ACCENT, side="bottom", size_eighths=12)

p = doc.add_paragraph()
paragraph_spacing(p, after=10, line=1.4)
styled_run(p, "Hybrid Freemium  ·  Tiered Subscription  ·  Lifetime Escape Hatch", size=11.5, bold=True, color=BRAND)
p = doc.add_paragraph()
paragraph_spacing(p, after=10, line=1.4)
styled_run(p,
    "Free tier acquires the audience and is monetised by AdMob. The paid ladder is "
    "intentionally narrow — too many tiers cause analysis paralysis. Annual is positioned "
    "as the default; Lifetime exists to capture subscription-fatigued users without "
    "diluting MRR.", size=10, color=INK_2)

# Pricing table
tiers = [
    ("Free",            "$0",                 "Forever",       "Ad-supported. 3 playlists, 50 clips, local only, basic schedule, sleep mode.",                       BRAND_SOFT_2),
    ("Premium Monthly", "$4.99 / mo",         "Recurring",     "No ads, unlimited playlists, cloud sync, multi-device, higher recording quality, all content packs voucher (1/mo).", WHITE),
    ("Premium Annual",  "$39.99 / yr",        "33% saving",    "Same as Monthly, billed yearly. Default tier — push hardest in onboarding.",                          ACCENT_SOFT),
    ("Lifetime",        "$79.99",             "One-time",      "All Premium features forever. Launch-window discount: $59.99 in the first 90 days.",                  WHITE),
    ("Family",          "$7.99 / mo",         "Up to 6",       "Premium for up to 6 family members. Increases ARPU per household, lowers blended churn.",             WHITE),
    ("Coach (B2B)",     "$24.99 / mo",        "Pro account",   "Premium + client-share, professional badge, branded share links. Bulk seats negotiable.",             WHITE),
]

tbl = doc.add_table(rows=len(tiers) + 1, cols=4)
tbl.autofit = False
widths = (Cm(3.6), Cm(2.6), Cm(2.6), Cm(8.6))
for col, w in zip(tbl.columns, widths):
    col.width = w

# Header row
hdrs = ("TIER", "PRICE", "MODE", "WHAT'S INCLUDED")
for i, h in enumerate(hdrs):
    c = tbl.rows[0].cells[i]
    c.width = widths[i]
    set_cell_shading(c, BRAND_DEEP)
    set_cell_borders(c, color=BRAND_DEEP, size_eighths=2)
    set_cell_margins(c, top=80, left=120, bottom=80, right=120)
    p = c.paragraphs[0]
    paragraph_spacing(p, after=0)
    styled_run(p, h, size=8.5, bold=True, color=ACCENT)

# Data rows
for r, (name, price, mode, desc, bg) in enumerate(tiers, start=1):
    row = tbl.rows[r]
    for ci, val in enumerate((name, price, mode, desc)):
        c = row.cells[ci]
        c.width = widths[ci]
        set_cell_shading(c, bg)
        set_cell_borders(c, color=LINE, size_eighths=2)
        set_cell_margins(c, top=90, left=120, bottom=90, right=120)
        p = c.paragraphs[0]
        paragraph_spacing(p, after=0, line=1.3)
        if ci == 0:
            styled_run(p, val, size=10, bold=True, color=BRAND_DEEP)
        elif ci == 1:
            styled_run(p, val, size=10, bold=True, color=INK)
        elif ci == 2:
            styled_run(p, val, size=8.5, bold=True, color=ACCENT if bg == ACCENT_SOFT else MUTED)
        else:
            styled_run(p, val, size=9, color=INK_2)

# Spacer
sp = doc.add_paragraph()
paragraph_spacing(sp, before=2, after=4)

# Why this pricing — two-column rationale
p = doc.add_paragraph()
paragraph_spacing(p, before=4, after=4)
styled_run(p, "Why This Structure Works", size=11.5, bold=True, color=BRAND_DEEP)

reasons = [
    ("Anchored on Annual", "Annual at $39.99 vs Monthly at $4.99 × 12 = $59.88 makes the saving obvious. 60–70% of paying users pick annual when framed this way."),
    ("Lifetime as guard-rail", "Captures the 5–10% who refuse subscriptions outright. Higher cash up-front and these users become the most loyal advocates."),
    ("Free trial drives signups", "7-day free Premium trial converts at 3–5%; day-5 auto-renew warning keeps refund noise low and store reviews healthy."),
    ("Regional pricing", "Set USD as anchor; let Apple / Google apply regional rules so PKR / INR / AED users see ~40–60% prices and don't bounce on cost."),
]
for label, body in reasons:
    p = doc.add_paragraph()
    paragraph_spacing(p, after=3, line=1.35)
    styled_run(p, "●  ", size=9.5, bold=True, color=ACCENT)
    styled_run(p, f"{label}.  ", size=9.5, bold=True, color=BRAND_DEEP)
    styled_run(p, body, size=9.5, color=INK_2)


# =====================================================================
# PAGE 4 — 90-DAY GTM PLAYBOOK + CLOSING
# =====================================================================
add_page_break(doc)

p = doc.add_paragraph()
paragraph_spacing(p, before=0, after=2)
styled_run(p, "03", size=9, bold=True, color=ACCENT)
p = doc.add_paragraph()
paragraph_spacing(p, before=0, after=4)
styled_run(p, "90-Day GTM Playbook", size=16, bold=True, color=BRAND_DEEP)
p = doc.add_paragraph()
paragraph_spacing(p, before=0, after=10)
paragraph_border(p, hex_color=ACCENT, side="bottom", size_eighths=12)

phases = [
    ("DAYS 0–30", "Launch & Acquisition",
     [
        "Soft launch in Pakistan + GCC — highest cultural fit for Urdu / Arabic content.",
        "Free tier live with 7-day Premium trial; AdMob enabled; Lifetime at $59.99 launch price.",
        "Seed 6–8 free curated packs (EN + UR + AR) and run ASO around “affirmations”, “whisper”, “sleep audio”.",
     ]),
    ("DAYS 31–60", "Convert & Expand",
     [
        "Paid acquisition on Meta + TikTok with PK + GCC creator partnerships.",
        "Ship first 3 paid content packs (celebrity / niche-creator collabs); open Coach (B2B) waitlist.",
        "Introduce Family plan; A/B test annual price points ($34.99 / $39.99 / $44.99).",
     ]),
    ("DAYS 61–90", "Compound & Refine",
     [
        "Launch Coach Pro tier publicly with 30 onboarded professionals; start affiliate program.",
        "Geo-expand to India + global English; localise pricing for INR / IDR / TRY.",
        "Lock funnel targets: ≥3% trial→paid, ≥40% annual mix, <5% monthly churn.",
     ]),
]

for tag, name, items in phases:
    # Heading row
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Cm(3.6)
    t.columns[1].width = Cm(13.8)

    cl = t.rows[0].cells[0]
    set_cell_shading(cl, ACCENT)
    set_cell_borders(cl, color=ACCENT, size_eighths=2)
    set_cell_margins(cl, top=80, left=140, bottom=80, right=140)
    p = cl.paragraphs[0]
    paragraph_spacing(p, after=0)
    styled_run(p, tag, size=8.5, bold=True, color=BRAND_DEEP)

    cr = t.rows[0].cells[1]
    set_cell_shading(cr, BRAND_SOFT)
    set_cell_borders(cr, color=BRAND_SOFT, size_eighths=2)
    set_cell_margins(cr, top=80, left=140, bottom=80, right=140)
    p = cr.paragraphs[0]
    paragraph_spacing(p, after=0)
    styled_run(p, name, size=11, bold=True, color=BRAND_DEEP)

    # Items
    for it in items:
        p = doc.add_paragraph()
        paragraph_spacing(p, before=1, after=1, line=1.3)
        pf = p.paragraph_format
        pf.left_indent = Cm(0.5)
        styled_run(p, "▸  ", size=9.5, bold=True, color=ACCENT)
        styled_run(p, it, size=9.5, color=INK_2)

    sp = doc.add_paragraph()
    paragraph_spacing(sp, before=0, after=1)

# KPI strip
p = doc.add_paragraph()
paragraph_spacing(p, before=4, after=4)
styled_run(p, "Success Metrics — Day 90", size=11.5, bold=True, color=BRAND_DEEP)

kpi_tbl = doc.add_table(rows=1, cols=4)
kpi_tbl.autofit = False
for col in kpi_tbl.columns:
    col.width = Cm(4.35)
kpis = [
    ("Trial → Paid",  "≥ 3%"),
    ("Annual Mix",    "≥ 40%"),
    ("Monthly Churn", "< 5%"),
    ("Day-30 ARPU",   "$0.40+"),
]
for i, (k, v) in enumerate(kpis):
    c = kpi_tbl.rows[0].cells[i]
    c.width = Cm(4.35)
    set_cell_shading(c, BRAND_DEEP)
    set_cell_borders(c, color=BRAND_DEEP, size_eighths=2)
    set_cell_margins(c, top=120, left=140, bottom=120, right=140)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_spacing(p, after=2)
    styled_run(p, k, size=8, bold=True, color=ACCENT)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_spacing(p2, after=0)
    styled_run(p2, v, size=14, bold=True, color=WHITE)

# Closing
p = doc.add_paragraph()
paragraph_spacing(p, before=10, after=4, line=1.4)
styled_run(p,
    "The model is deliberately conservative on price and aggressive on packaging. "
    "It lets WhisperBack acquire users at the lowest possible friction, monetise the "
    "free majority through ads, convert the engaged minority through a clean subscription "
    "ladder, and unlock long-tail revenue through content packs and B2B — all without "
    "asking the team to build five products.",
    size=10, color=INK_2)

# Footer band
foot = doc.add_table(rows=1, cols=2)
foot.autofit = False
foot.columns[0].width = Cm(11)
foot.columns[1].width = Cm(6.4)
fl = foot.rows[0].cells[0]
fr = foot.rows[0].cells[1]
for c in (fl, fr):
    set_cell_shading(c, BRAND_DEEP)
    set_cell_borders(c, color=BRAND_DEEP, size_eighths=2)
    set_cell_margins(c, top=100, left=200, bottom=100, right=200)
p = fl.paragraphs[0]
paragraph_spacing(p, after=0)
styled_run(p, "Prepared by FOUS Ventures  ·  Strategy & Discovery", size=8.5, bold=True, color=WHITE)
p = fr.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
paragraph_spacing(p, after=0)
styled_run(p, "WhisperBack  ·  Monetization v1.0", size=8.5, color=ACCENT)


# Save
out = "WhisperBack_Monetization_Strategy.docx"
doc.save(out)
print(f"Saved: {out}")
